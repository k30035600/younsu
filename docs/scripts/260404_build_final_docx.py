# -*- coding: utf-8 -*-
"""행정심판청구(최종)/260404 Markdown(01·02)을 공문서 스타일에 가깝게 DOCX로 변환한다.

- 글꼴: **나눔명조**. 제출 본문 단락 정렬은 **양쪽 맞춤**(공문 일반 형식). 증거 목록은 **표**로 넣어 경로·파일명 줄의 자간 과다를 완화함. 말미 **【판례 각주】** 목록은 **왼쪽**·들여쓰기(번호 대응 가독성).
- **청구·신청 원문**: **파란색**, 본문 **12pt**, `##` **14pt**, `###` **12pt**, 표제 **24pt**.
- **제출 원문**: 파란색, 본문 12pt 이상(표제·소제목 규칙 동일).
- **작성자 검토·검수·주의(제출 전 삭제 권장)**: 줄 말머리 **`[작업] `·`[검토] `·`[검수] `·`[주의] `·`[편집] `** 는 **검정 10pt**, 접두 **그대로 Word에 표시**, 왼쪽 들여쓰기·앞여백으로 본문과 구분. 인용 `> `, `## [참고]…` 이하·말미 판례 각주도 **검정 10pt**.
- 본문(파란) 안 `*…*` 는 **검정 10pt**(비이탤릭)로 원문과 구분.
- 인용 판례(사건번호 패턴) 뒤에 **위첨자 각주 번호**를 붙이고, 문서 말미(새 쪽)에
  **【판례 각주】** 목록을 둔다. (python-docx는 Word 네이티브 각주 부분이 없어
  말미 목록으로 두며, Word에서 필요 시 「참고 각주」로 변환 가능.)

- 입력: `행정심판청구(최종)/{yymmdd}_01_…_최종.md`, `…/{yymmdd}_02_…_최종.md`
- 출력 DOCX: **`docs/{yymmdd}/`** — 접두 6자리(`260404` 등)는 **01번 MD 파일명**에서 자동 추출. `--yymmdd` 로 덮어쓸 수 있음.
- `**[증거자료 목록]**` ~ `**붙임**` 직전: 안내 문단은 본문 단락, `N.`·`N-M.` 항목은 표(표 스타일 Table Grid).
- 여백·줄띄움·붙임·끝. 표기는 `행정심판청구(최종)/2025 개정 공문서 작성법_박종덕.pdf`(동일 내용
  `행정심판청구(최종)/{yymmdd}_참고_공문서작성법_PDF텍스트추출.txt`)의 취지에 맞춤(행정 공문 일반 서식에 가까운 A4 여백).

의존성: pip install python-docx pypdf

실행(프로젝트 루트):
  python docs/scripts/260404_build_final_docx.py
  python docs/scripts/260404_build_final_docx.py --extract-pdf
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor
from docx.oxml.ns import qn

EVIDENCE_HEAD = "**[증거자료 목록]**"
_EVID_ITEM = re.compile(r"^(\d+(?:-\d+)?)\.\s+(.+)$")
# 포털용 MD 앵커 `[갑 제N호증](#N)` 은 Word 제출본에서는 대괄호·(#…) 없이 표시만 남긴다.
_MD_INTERNAL_ANCHOR = re.compile(r"\[([^\]]*)\]\(#[^)]+\)")


def _strip_internal_md_anchors(text: str) -> str:
    return _MD_INTERNAL_ANCHOR.sub(r"\1", text)

_REPO = Path(__file__).resolve().parents[2]
_FINAL_ROOT = _REPO / "행정심판청구(최종)"
PDF_STYLE = _FINAL_ROOT / "2025 개정 공문서 작성법_박종덕.pdf"

MD_01 = _FINAL_ROOT / "260404_01_행정심판청구서_최종.md"
MD_02 = _FINAL_ROOT / "260404_02_집행정지신청서_최종.md"

_YYMMDD_FROM_STEM = re.compile(r"^(\d{6})_")


def _docs_out_dir(md01: Path, yymmdd_override: str | None) -> Path:
    """DOCX·산출물은 항상 `docs/{yymmdd}/`."""
    if yymmdd_override:
        return _REPO / "docs" / yymmdd_override.strip()
    m = _YYMMDD_FROM_STEM.match(md01.name)
    if m:
        return _REPO / "docs" / m.group(1)
    return _REPO / "docs" / "260404"


def _out_extract_txt(yymmdd: str) -> Path:
    return _FINAL_ROOT / f"{yymmdd}_참고_공문서작성법_PDF텍스트추출.txt"


def _yymmdd_str(md01: Path, yymmdd_override: str | None) -> str:
    if yymmdd_override:
        return yymmdd_override.strip()
    m = _YYMMDD_FROM_STEM.match(md01.name)
    return m.group(1) if m else "260404"


# Windows 한글판 일반 명칭. 없으면 Word 기본 글꼴로 대체됨.
FONT = "나눔명조"
RGB_BLUE = RGBColor(0x00, 0x33, 0xCC)
RGB_BLACK = RGBColor(0x00, 0x00, 0x00)
# 공문 본문·소제목: 양쪽 맞춤. 표 머리·말미 각주만 예외.
ALIGN_BODY = WD_ALIGN_PARAGRAPH.JUSTIFY
ALIGN_TABLE_HEADER = WD_ALIGN_PARAGRAPH.CENTER
ALIGN_FOOTNOTE_BLOCK = WD_ALIGN_PARAGRAPH.LEFT
# 제출 원문이 아닌 «작성자용» 줄 — 길이 내림차순으로 매칭(더 긴 접두 우선).
NOTE_LINE_PREFIXES: tuple[str, ...] = (
    "[편집] ",
    "[검수] ",
    "[검토] ",
    "[주의] ",
    "[작업] ",
)
TASK_PREFIX = "[작업] "
REF_HEAD = "## [참고]"
NOTE_LEFT_INDENT_MM = 5.0
NOTE_GAP_BEFORE_PT = 10.0
MAIN_GAP_AFTER_NOTE_PT = 8.0

# 본문(파란) vs 메모(검정) 글자 크기
SZ_MAIN_BODY = 12
SZ_MAIN_H2 = 14
SZ_MAIN_H3 = 12
SZ_TITLE = 24
SZ_NOTE = 10
# 본문(파란) 안 *…* 주석
SZ_INLINE_NOTE = 10

_BOLD_SPLIT = re.compile(r"(\*\*[^*]+\*\*)")
# 대법원 형식 사건번호(예: 2008두167, 91누13441, 99다70600)
_CASE_ID = re.compile(r"\d{4}두\d+|\d{2}누\d+|\d{2}다\d+")
_SUP_TRANS = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")


def _superscript_num(n: int) -> str:
    return str(n).translate(_SUP_TRANS)


class FootnoteState:
    """사건번호별 동일 번호 재사용."""

    def __init__(self) -> None:
        self._case_to_num: dict[str, int] = {}
        self.num_to_note: dict[int, str] = {}
        self._next = 1

    def mark(self, case_id: str) -> int:
        if case_id in self._case_to_num:
            return self._case_to_num[case_id]
        n = self._next
        self._next += 1
        self._case_to_num[case_id] = n
        self.num_to_note[n] = (
            f"{case_id} — 국가법령정보센터(law.go.kr)에서 동일 사건번호로 조회·원문 대조."
        )
        return n


def _append_runs_case_ids(
    p,
    seg: str,
    fn: FootnoteState | None,
    *,
    size: int,
    color_rgb: RGBColor,
    bold_base: bool,
    sup_floor: int,
) -> None:
    seg = seg.replace("`", "")
    if not seg:
        return
    if fn is None:
        run = p.add_run(seg)
        set_font(run, size=size, color_rgb=color_rgb)
        run.bold = bold_base
        return
    pos = 0
    for m in _CASE_ID.finditer(seg):
        if m.start() > pos:
            run = p.add_run(seg[pos : m.start()])
            set_font(run, size=size, color_rgb=color_rgb)
            run.bold = bold_base
        case = m.group()
        run_case = p.add_run(case)
        set_font(run_case, size=size, color_rgb=color_rgb)
        run_case.bold = bold_base
        fid = fn.mark(case)
        run_sup = p.add_run(_superscript_num(fid))
        set_font(run_sup, size=max(size - 1, sup_floor), color_rgb=color_rgb)
        run_sup.bold = False
        run_sup.font.superscript = True
        pos = m.end()
    if pos < len(seg):
        run = p.add_run(seg[pos:])
        set_font(run, size=size, color_rgb=color_rgb)
        run.bold = bold_base


def _append_runs_with_inline_notes(
    p,
    seg: str,
    fn: FootnoteState | None,
    *,
    main_size: int,
    color_rgb: RGBColor,
    bold_base: bool,
    sup_floor: int,
    use_inline_note_style: bool,
) -> None:
    """`*한 줄 주석*` 은 파란 본문에서만 검정 10pt(비이탤릭)."""
    if not use_inline_note_style:
        _append_runs_case_ids(
            p,
            seg,
            fn,
            size=main_size,
            color_rgb=color_rgb,
            bold_base=bold_base,
            sup_floor=sup_floor,
        )
        return
    parts = re.split(r"(\*[^*]+\*)", seg)
    for chunk in parts:
        if not chunk:
            continue
        if chunk.startswith("*") and chunk.endswith("*") and len(chunk) >= 2:
            inner = chunk[1:-1].replace("`", "")
            if not inner:
                continue
            run = p.add_run(inner)
            set_font(run, size=SZ_INLINE_NOTE, color_rgb=RGB_BLACK)
            run.italic = False
            run.bold = False
            continue
        _append_runs_case_ids(
            p,
            chunk,
            fn,
            size=main_size,
            color_rgb=color_rgb,
            bold_base=bold_base,
            sup_floor=sup_floor,
        )


def set_font(
    run,
    font_name: str = FONT,
    size: int = 12,
    color_rgb: RGBColor | None = None,
) -> None:
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:hint"), "eastAsia")
    run.font.size = Pt(size)
    if color_rgb is not None:
        run.font.color.rgb = color_rgb


def _page_margins(doc: Document) -> None:
    """2025 공문서 작성법 길라잡이(박종덕) 실무 서식에 가깝게: 본문 가독성·제본 여유."""
    for section in doc.sections:
        section.top_margin = Mm(32)
        section.bottom_margin = Mm(22)
        section.left_margin = Mm(28)
        section.right_margin = Mm(26)


def _add_mixed_paragraph(
    doc: Document,
    text: str,
    fn: FootnoteState | None,
    *,
    size: int = 12,
    bold_default: bool = False,
    indent_mm: float = 0,
    align=ALIGN_BODY,
    line_spacing: float = 1.65,
    color_rgb: RGBColor = RGB_BLUE,
    space_after_pt: float | None = None,
    space_before_pt: float | None = None,
) -> None:
    text = _strip_internal_md_anchors(text.rstrip())
    if not text:
        doc.add_paragraph()
        return
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    if space_before_pt is not None:
        p.paragraph_format.space_before = Pt(space_before_pt)
    if space_after_pt is None:
        space_after_pt = 5.0 if size <= 9 else 8.0
    p.paragraph_format.space_after = Pt(space_after_pt)
    if indent_mm:
        p.paragraph_format.left_indent = Mm(indent_mm)

    sup_floor = 7 if size <= 9 else 9
    use_inline_note_style = color_rgb == RGB_BLUE and size >= SZ_MAIN_BODY

    for part in _BOLD_SPLIT.split(text):
        if not part:
            continue
        is_bold_seg = part.startswith("**") and part.endswith("**")
        seg = part[2:-2] if is_bold_seg else part
        base_bold = bold_default or is_bold_seg
        _append_runs_with_inline_notes(
            p,
            seg,
            fn,
            main_size=size,
            color_rgb=color_rgb,
            bold_base=base_bold,
            sup_floor=sup_floor,
            use_inline_note_style=use_inline_note_style,
        )


def _add_heading_center(
    doc: Document, text: str, size: int = 24, color_rgb: RGBColor = RGB_BLUE
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 표제와 «사 건» 사이 — 미리보기·목록 파싱과 무관하게 Word에서 간격이 나오도록 여백 보강
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(_strip_internal_md_anchors(text.strip()))
    set_font(run, size=size, color_rgb=color_rgb)
    run.bold = True


def _is_evidence_heading(stripped: str) -> bool:
    s = stripped.replace(" ", "")
    return s == EVIDENCE_HEAD.replace(" ", "") or stripped.strip() == "[증거자료 목록]"


def _consume_evidence_block(lines: list[str], start: int) -> tuple[list[str], int]:
    j = start
    out: list[str] = []
    while j < len(lines):
        st = lines[j].strip()
        if st.startswith("**붙임**"):
            break
        if re.fullmatch(r"\*\*끝\.\*\*", st):
            break
        if st.startswith("**청구인**"):
            break
        out.append(lines[j])
        j += 1
    return out, j


def _md_plain_cell(t: str) -> str:
    t = _strip_internal_md_anchors(t)
    t = t.replace("**", "").replace("`", "")
    t = re.sub(r"\s+", " ", t).strip()
    return t


def _parse_evidence_block(block_lines: list[str]) -> tuple[list[str], list[tuple[str, str, str]]]:
    intro_raw: list[str] = []
    rows: list[tuple[str, str, str]] = []
    phase = "intro"
    current: tuple[str, str, str] | None = None

    for line in block_lines:
        st = line.strip()
        if not st:
            if phase == "intro":
                intro_raw.append("")
            continue
        m = _EVID_ITEM.match(st)
        if m:
            phase = "items"
            if current:
                rows.append(current)
            num = m.group(1) + "."
            rest = m.group(2).strip()
            if " — " in rest:
                left, right = rest.split(" — ", 1)
            else:
                left, right = rest, ""
            current = (num, _md_plain_cell(left), _md_plain_cell(right))
        elif phase == "items" and current is not None:
            extra = _md_plain_cell(st)
            if extra:
                c0, c1, c2 = current
                current = (c0, c1, (c2 + " " if c2 else "") + extra)
        else:
            intro_raw.append(line)

    if current:
        rows.append(current)
    return intro_raw, rows


def _style_table_cell_text(
    cell,
    text: str,
    *,
    size: int,
    bold: bool,
    color_rgb: RGBColor,
    align: WD_ALIGN_PARAGRAPH = ALIGN_BODY,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text if text else " ")
    set_font(run, size=size, color_rgb=color_rgb)
    run.bold = bold


def _add_evidence_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows) + 1, cols=3)
    tbl.style = "Table Grid"
    headers = ("번호", "호증·취지(요지)", "파일·폴더(비고)")
    for ci, ht in enumerate(headers):
        _style_table_cell_text(
            tbl.rows[0].cells[ci],
            ht,
            size=SZ_MAIN_BODY,
            bold=True,
            color_rgb=RGB_BLUE,
            align=ALIGN_TABLE_HEADER,
        )
    for ri, (no, left, right) in enumerate(rows, start=1):
        cells = tbl.rows[ri].cells
        _style_table_cell_text(
            cells[0], no, size=SZ_MAIN_BODY, bold=False, color_rgb=RGB_BLUE
        )
        _style_table_cell_text(
            cells[1], left, size=SZ_MAIN_BODY, bold=False, color_rgb=RGB_BLUE
        )
        _style_table_cell_text(
            cells[2], right, size=SZ_MAIN_BODY, bold=False, color_rgb=RGB_BLUE
        )
    doc.add_paragraph()


def _append_footnote_pages(doc: Document, fn: FootnoteState) -> None:
    if not fn.num_to_note:
        return
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("【판례 각주】")
    set_font(r, size=SZ_NOTE, color_rgb=RGB_BLACK)
    r.bold = True
    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.alignment = ALIGN_FOOTNOTE_BLOCK
    intro.paragraph_format.line_spacing = 1.6
    intro.paragraph_format.space_after = Pt(6)
    t = intro.add_run(
        "본문 위첨자 번호와 대응합니다. 제출 전 law.go.kr 원문과 대조하십시오."
    )
    set_font(t, size=SZ_NOTE, color_rgb=RGB_BLACK)
    doc.add_paragraph()

    for n in sorted(fn.num_to_note):
        p2 = doc.add_paragraph()
        p2.alignment = ALIGN_FOOTNOTE_BLOCK
        p2.paragraph_format.line_spacing = 1.6
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.left_indent = Mm(5)
        p2.paragraph_format.first_line_indent = Mm(-5)
        mark = p2.add_run(f"{n}. ")
        set_font(mark, size=SZ_NOTE, color_rgb=RGB_BLACK)
        mark.bold = True
        body = p2.add_run(fn.num_to_note[n])
        set_font(body, size=SZ_NOTE, color_rgb=RGB_BLACK)


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    raw = md_path.read_text(encoding="utf-8")
    lines = raw.replace("\r\n", "\n").split("\n")

    doc = Document()
    _page_margins(doc)
    fn = FootnoteState()
    in_reference = False
    # 직전 실질 단락: 제출 원문(main) vs 작성자용(note) — 여백으로 구분
    last_block_kind: str = "neutral"
    # **끝.** 다음(빈 줄만 허용) 첫 **청구인** / **신청인**만 서명란(오른쪽 맞춤). «당사자» 절 첫 줄은 일반 본문.
    pending_signature_after_끝: bool = False

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        note_line = False
        for pfx in sorted(NOTE_LINE_PREFIXES, key=len, reverse=True):
            if s.startswith(pfx):
                note_line = True
                break

        if _is_evidence_heading(s):
            sb_h = MAIN_GAP_AFTER_NOTE_PT if last_block_kind == "note" else None
            _add_mixed_paragraph(
                doc,
                "【증거자료 목록】",
                fn,
                size=SZ_MAIN_H2,
                bold_default=True,
                indent_mm=0,
                color_rgb=RGB_BLUE,
                space_after_pt=10.0,
                space_before_pt=sb_h,
            )
            i += 1
            block_lines, i = _consume_evidence_block(lines, i)
            intro_raw, ev_rows = _parse_evidence_block(block_lines)
            for il in intro_raw:
                if not il.strip():
                    doc.add_paragraph()
                    continue
                _add_mixed_paragraph(
                    doc,
                    il.strip(),
                    fn,
                    size=SZ_MAIN_BODY,
                    indent_mm=0,
                    color_rgb=RGB_BLUE,
                )
            _add_evidence_table(doc, ev_rows)
            last_block_kind = "main"
            continue

        if not s:
            doc.add_paragraph()
            i += 1
            continue
        if s == "---":
            doc.add_paragraph()
            last_block_kind = "neutral"
            i += 1
            continue
        if s.startswith(REF_HEAD):
            in_reference = True

        use_task = note_line or in_reference
        rgb = RGB_BLACK if use_task else RGB_BLUE
        sz_body = SZ_NOTE if use_task else SZ_MAIN_BODY
        sz_h2 = SZ_NOTE if use_task else SZ_MAIN_H2
        sz_h3 = SZ_NOTE if use_task else SZ_MAIN_H3

        st = s
        if pending_signature_after_끝:
            if not (
                re.fullmatch(r"\*\*끝\.\*\*", st)
                or re.match(r"^\*\*청구인\*\*", st)
                or re.match(r"^\*\*신청인\*\*", st)
            ):
                pending_signature_after_끝 = False

        if re.fullmatch(r"\*\*끝\.\*\*", st):
            pending_signature_after_끝 = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.65
            ru = p.add_run("끝.")
            set_font(ru, size=SZ_MAIN_BODY, color_rgb=RGB_BLUE)
            last_block_kind = "main"
            i += 1
            continue

        m_sig = re.match(r"^\*\*청구인\*\*\s*(.*)$", st)
        if m_sig:
            if pending_signature_after_끝:
                pending_signature_after_끝 = False
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(14)
                p.paragraph_format.line_spacing = 1.65
                r0 = p.add_run("청구인")
                set_font(r0, size=SZ_MAIN_BODY, color_rgb=RGB_BLUE)
                r0.bold = True
                tail = m_sig.group(1).strip()
                if tail:
                    r1 = p.add_run("  " + tail)
                    set_font(r1, size=SZ_MAIN_BODY, color_rgb=RGB_BLUE)
                last_block_kind = "main"
                i += 1
                continue
            sb_sig = MAIN_GAP_AFTER_NOTE_PT if last_block_kind == "note" else None
            _add_mixed_paragraph(
                doc,
                st,
                fn,
                size=sz_body,
                indent_mm=0,
                color_rgb=rgb,
                space_before_pt=sb_sig,
            )
            last_block_kind = "main" if not in_reference else "note"
            i += 1
            continue

        m_app = re.match(r"^\*\*신청인\*\*\s*(.*)$", st)
        if m_app:
            if pending_signature_after_끝:
                pending_signature_after_끝 = False
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(14)
                p.paragraph_format.line_spacing = 1.65
                r0 = p.add_run("신청인")
                set_font(r0, size=SZ_MAIN_BODY, color_rgb=RGB_BLUE)
                r0.bold = True
                tail = m_app.group(1).strip()
                if tail:
                    r1 = p.add_run("  " + tail)
                    set_font(r1, size=SZ_MAIN_BODY, color_rgb=RGB_BLUE)
                last_block_kind = "main"
                i += 1
                continue
            sb_app = MAIN_GAP_AFTER_NOTE_PT if last_block_kind == "note" else None
            _add_mixed_paragraph(
                doc,
                st,
                fn,
                size=sz_body,
                indent_mm=0,
                color_rgb=rgb,
                space_before_pt=sb_app,
            )
            last_block_kind = "main" if not in_reference else "note"
            i += 1
            continue

        if s.startswith("# ") and not s.startswith("## "):
            title = s[2:].strip()
            title = re.sub(r"^\[(?:최종본|행정심판최종본)\]\s*", "", title)
            _add_heading_center(doc, title, SZ_TITLE, color_rgb=rgb)
            last_block_kind = "main"
            i += 1
            continue
        if s.startswith("## ") and not s.startswith("### "):
            sb_h2 = None
            if in_reference and last_block_kind == "main":
                sb_h2 = NOTE_GAP_BEFORE_PT
            elif not in_reference and last_block_kind == "note":
                sb_h2 = MAIN_GAP_AFTER_NOTE_PT
            _add_mixed_paragraph(
                doc,
                s[3:].strip(),
                fn,
                size=sz_h2,
                bold_default=True,
                indent_mm=0,
                color_rgb=rgb,
                space_after_pt=10.0,
                space_before_pt=sb_h2,
            )
            last_block_kind = "note" if in_reference else "main"
            i += 1
            continue
        if s.startswith("### ") and not s.startswith("#### "):
            sb_h3 = MAIN_GAP_AFTER_NOTE_PT if (not in_reference and last_block_kind == "note") else None
            _add_mixed_paragraph(
                doc,
                s[4:].strip(),
                fn,
                size=sz_h3,
                bold_default=True,
                indent_mm=5,
                color_rgb=rgb,
                space_before_pt=sb_h3,
            )
            last_block_kind = "note" if in_reference else "main"
            i += 1
            continue
        if s.startswith("#### "):
            sb_h4 = MAIN_GAP_AFTER_NOTE_PT if (not in_reference and last_block_kind == "note") else None
            _add_mixed_paragraph(
                doc,
                s[5:].strip(),
                fn,
                size=sz_h3,
                bold_default=True,
                indent_mm=8,
                color_rgb=rgb,
                space_before_pt=sb_h4,
            )
            last_block_kind = "note" if in_reference else "main"
            i += 1
            continue
        if s.startswith("> "):
            sb_q = NOTE_GAP_BEFORE_PT if last_block_kind == "main" else None
            _add_mixed_paragraph(
                doc,
                s[2:].strip(),
                fn,
                size=SZ_NOTE,
                indent_mm=max(10.0, NOTE_LEFT_INDENT_MM),
                color_rgb=RGB_BLACK,
                space_before_pt=sb_q,
            )
            last_block_kind = "note"
            i += 1
            continue
        if note_line:
            sb_n = NOTE_GAP_BEFORE_PT if last_block_kind == "main" else None
            _add_mixed_paragraph(
                doc,
                s,
                fn,
                size=SZ_NOTE,
                indent_mm=NOTE_LEFT_INDENT_MM,
                color_rgb=RGB_BLACK,
                space_before_pt=sb_n,
            )
            last_block_kind = "note"
            i += 1
            continue
        sb_body = MAIN_GAP_AFTER_NOTE_PT if last_block_kind == "note" else None
        _add_mixed_paragraph(
            doc,
            s,
            fn,
            size=sz_body,
            indent_mm=0,
            color_rgb=rgb,
            space_before_pt=sb_body,
        )
        last_block_kind = "main" if not in_reference else "note"

        i += 1

    _append_footnote_pages(doc, fn)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print("작성:", docx_path.relative_to(_REPO), f"(판례 각주 {len(fn.num_to_note)}건)")


def extract_pdf_text(pdf_path: Path, out_txt: Path) -> None:
    try:
        from pypdf import PdfReader
    except ImportError:
        print("pypdf 필요: pip install pypdf", file=sys.stderr)
        raise
    if not pdf_path.is_file():
        print("경고: PDF 없음 —", pdf_path, file=sys.stderr)
        return
    reader = PdfReader(str(pdf_path))
    chunks: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception as e:
            t = f"[페이지 {i + 1} 추출 오류: {e}]"
        chunks.append(f"--- 페이지 {i + 1} ---\n{t}\n")
    out_txt.write_text("\n".join(chunks), encoding="utf-8")
    print("추출:", out_txt.relative_to(_REPO), f"({len(reader.pages)}쪽)")


def main() -> None:
    ap = argparse.ArgumentParser(
        description="행정심판청구(최종)/*_최종.md → docs/{yymmdd}/ docx (yymmdd는 01번 MD 파일명 접두 또는 --yymmdd)"
    )
    ap.add_argument("--01", dest="only01", action="store_true", help="청구서만")
    ap.add_argument("--02", dest="only02", action="store_true", help="집행정지만")
    ap.add_argument(
        "--md01",
        type=Path,
        default=None,
        help=f"청구서 MD (기본: {MD_01.name})",
    )
    ap.add_argument(
        "--md02",
        type=Path,
        default=None,
        help=f"집행정지 MD (기본: {MD_02.name})",
    )
    ap.add_argument(
        "--yymmdd",
        default=None,
        metavar="YYMMDD",
        help="출력 폴더 docs/YYMMDD/ (미지정 시 --md01 파일명의 6자리 접두)",
    )
    ap.add_argument(
        "--extract-pdf",
        action="store_true",
        help="공문서 PDF 텍스트만 추출 → 행정심판청구(최종)/{yymmdd}_참고_공문서작성법_PDF텍스트추출.txt",
    )
    ap.add_argument(
        "--pdf",
        type=Path,
        default=PDF_STYLE,
        help="공문서 작성법 PDF 경로",
    )
    args = ap.parse_args()

    md01 = args.md01 if args.md01 is not None else MD_01
    md02 = args.md02 if args.md02 is not None else MD_02
    yymmdd = _yymmdd_str(md01, args.yymmdd)
    out_extract = _out_extract_txt(yymmdd)
    docs_out = _docs_out_dir(md01, args.yymmdd)
    docx_01 = docs_out / f"{md01.stem}.docx"
    docx_02 = docs_out / f"{md02.stem}.docx"

    if args.extract_pdf:
        extract_pdf_text(args.pdf, out_extract)
        if not args.only01 and not args.only02:
            return

    do01 = args.only01 or (not args.only01 and not args.only02)
    do02 = args.only02 or (not args.only01 and not args.only02)

    try:
        import docx  # noqa: F401
    except ImportError:
        print("python-docx 필요: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    docs_out.mkdir(parents=True, exist_ok=True)

    if do01:
        if not md01.is_file():
            print("경고: 없음", md01, file=sys.stderr)
        else:
            md_to_docx(md01, docx_01)
    if do02:
        if not md02.is_file():
            print("경고: 없음", md02, file=sys.stderr)
        else:
            md_to_docx(md02, docx_02)


if __name__ == "__main__":
    main()
