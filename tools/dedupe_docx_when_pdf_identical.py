# -*- coding: utf-8 -*-
"""
같은 폴더에 `stem.docx`와 `stem.pdf`가 있을 때, 추출·정규화한 본문 텍스트가
동일하면 DOCX만 제거해 단일본(PDF)으로 둡니다.

  python tools/dedupe_docx_when_pdf_identical.py                    # 미리보기(완전 일치만)
  python tools/dedupe_docx_when_pdf_identical.py --min-ratio 0.9965 # 유사도 미리보기
  python tools/dedupe_docx_when_pdf_identical.py --apply --min-ratio 0.9965

기본(--min-ratio 1.0)은 추출 텍스트 문자열이 완전히 같을 때만 삭제합니다. PDF는 줄바꿈·표 순서
등으로 Word와 추출문이 달라지는 경우가 많아, 같은 원고에서 낸 PDF라면 `--min-ratio`로
difflib.SequenceMatcher 비율(0~1) 임계값을 낮출 수 있습니다(낮출수록 오판 위험 증가).
"""
from __future__ import annotations

import argparse
import difflib
import re
import sys
import unicodedata
from pathlib import Path

_REPO = Path(__file__).resolve().parent.parent

try:
    from docx import Document
except ImportError:
    print("pip install python-docx", file=sys.stderr)
    raise
try:
    from pypdf import PdfReader
except ImportError:
    print("pip install pypdf", file=sys.stderr)
    raise

_SKIP_DIR_NAMES = {".git", "node_modules", "__pycache__", ".venv", "venv"}


def _rel_to_root(p: Path, root: Path) -> Path:
    try:
        return p.relative_to(root)
    except ValueError:
        return p


def _norm_text(s: str) -> str:
    s = unicodedata.normalize("NFKC", s or "")
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def _text_from_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(t)
    return _norm_text("\n".join(parts))


def _text_from_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            chunks.append(t)
    return _norm_text("\n".join(chunks))


def _iter_docx_with_pdf_sibling(root: Path) -> list[tuple[Path, Path]]:
    out: list[tuple[Path, Path]] = []
    for p in root.rglob("*.docx"):
        if any(part in _SKIP_DIR_NAMES for part in p.parts):
            continue
        pdf = p.with_suffix(".pdf")
        if pdf.is_file():
            out.append((p, pdf))
    return sorted(out, key=lambda x: str(x[0]).lower())


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--root",
        type=Path,
        default=_REPO,
        help="검색 루트(기본: 저장소 루트)",
    )
    ap.add_argument(
        "--apply",
        action="store_true",
        help="동일 판정 시 DOCX 삭제(미지정이면 목록만)",
    )
    ap.add_argument(
        "--min-ratio",
        type=float,
        default=1.0,
        metavar="R",
        help="1.0=추출문 완전 일치만. 1 미만이면 SequenceMatcher 비율>=R일 때 동일 처리(오판 주의)",
    )
    args = ap.parse_args()
    root = args.root.resolve()
    if not root.is_dir():
        raise SystemExit(f"not a directory: {root}")

    min_r = float(args.min_ratio)
    if not 0.0 < min_r <= 1.0:
        raise SystemExit("--min-ratio must be in (0, 1]")

    pairs = _iter_docx_with_pdf_sibling(root)
    same: list[tuple[Path, Path, str]] = []
    diff: list[tuple[Path, str]] = []

    for docx_path, pdf_path in pairs:
        try:
            td = _text_from_docx(docx_path)
            tp = _text_from_pdf(pdf_path)
        except Exception as e:
            diff.append((docx_path, f"read error: {e}"))
            continue
        if not td and not tp:
            ratio = 1.0
            match = True
        elif td == tp:
            ratio = 1.0
            match = True
        else:
            ratio = difflib.SequenceMatcher(None, td, tp).ratio()
            match = min_r < 1.0 and ratio >= min_r
        if match:
            tag = "exact" if ratio >= 1.0 - 1e-15 else f"ratio={ratio:.5f}"
            same.append((docx_path, pdf_path, tag))
        else:
            hint = f"len docx={len(td)} pdf={len(tp)} ratio={ratio:.5f}"
            if len(td) != len(tp):
                hint += f" (diff {abs(len(td) - len(tp))} chars)"
            diff.append((docx_path, hint))

    for d, _pdf, tag in same:
        rel = _rel_to_root(d, root)
        print(f"SAME\t{rel}\t{tag}")
        if args.apply:
            d.unlink()
            print(f"  removed\t{rel}")

    for d, msg in diff:
        rel = _rel_to_root(d, root)
        print(f"DIFF\t{rel}\t{msg}")

    print(f"\nSummary: {len(same)} same, {len(diff)} diff/skipped, pairs scanned={len(pairs)}")


if __name__ == "__main__":
    main()

