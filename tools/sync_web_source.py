# -*- coding: utf-8 -*-
"""`commission-portal/public/source` 동기화: 정본 MD 복사, site-display.json, PDF 생성.

- **포털 탭 본문**은 저장소 정본 `.md`를 `/serve/`로 읽습니다.
- **`public/source/*.md`**는 정본에서 복사한 **편집·보관·대조용** 사본이며, `/source/…`로도 서빙됩니다. (탭: 청구서·갑1~3·집행정지)
- **`.pdf`**는 같은 내용의 **제출·인쇄 보조**용으로 생성합니다(내부적으로 DOCX로 조판 후 변환).

실행(프로젝트 루트):
  python tools/build_commission_evidence_json.py
  python tools/survey_gab_evidence_full.py   # 선택
  python tools/sync_web_source.py

`npm start` 시 `start.js`가 `/source/…` 아래 `public/source` 파일을 서빙합니다.

의존성: `pip install python-docx`(MD→DOCX), DOCX→PDF는 Word COM·docx2pdf 또는 LibreOffice(`pdf_convert_util`).
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import tempfile
from pathlib import Path

from pdf_convert_util import convert_docx_to_pdf

_REPO = Path(__file__).resolve().parent.parent
_PORTAL_DATA = _REPO / "web" / "commission-portal" / "public" / "data" / "portal-data.json"
_PUBLIC_SOURCE = _REPO / "web" / "commission-portal" / "public" / "source"

# tabSources 키 → 저장 파일명(utf-8)
_TAB_FILES: list[tuple[str, str]] = [
    ("appeal", "행정심판청구.md"),
    ("gab1", "별지_갑1호증.md"),
    ("gab2", "별지_갑2호증.md"),
    ("gab3", "별지_갑3호증.md"),
    ("injunction", "집행정지신청.md"),
]

_SITE_DISPLAY_DEFAULT = {
    "siteTitle": "농원근린공원 행정심판청구",
    "siteSubtitle": "집행정지신청 병합 · 인천광역시 행정심판위원회 심리 참고",
    "updated": "2026-04-05",
}


def _repo_file(rel: str) -> Path:
    r = rel.replace("\\", "/").strip().lstrip("/")
    return _REPO / r


def _copy_tab_sources(tab_sources: dict) -> None:
    _PUBLIC_SOURCE.mkdir(parents=True, exist_ok=True)
    for key, dest_name in _TAB_FILES:
        rel = (tab_sources or {}).get(key)
        if not rel:
            print(f"skip {key}: tabSources에 없음", file=sys.stderr)
            continue
        src = _repo_file(str(rel))
        if not src.is_file():
            print(f"skip {key}: 파일 없음 {src}", file=sys.stderr)
            continue
        dest = _PUBLIC_SOURCE / dest_name
        shutil.copy2(src, dest)
        print(f"복사 {dest_name} ← {rel}")


def _write_site_display(meta: dict, force: bool) -> Path:
    """`portal-data.json`의 meta와 동일한 제목·부제·기준일을 site-display.json에 기록."""
    path = _PUBLIC_SOURCE / "site-display.json"
    _PUBLIC_SOURCE.mkdir(parents=True, exist_ok=True)
    m = meta or {}
    out = {
        "siteTitle": m.get("siteTitle") or _SITE_DISPLAY_DEFAULT["siteTitle"],
        "siteSubtitle": m.get("siteSubtitle") or _SITE_DISPLAY_DEFAULT["siteSubtitle"],
        "updated": m.get("updated") or _SITE_DISPLAY_DEFAULT["updated"],
    }
    if path.is_file() and not force:
        try:
            cur = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            cur = {}
        if cur == out:
            print(f"유지 {path.name} (portal meta와 동일)")
            return path
    path.write_text(json.dumps(out, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"작성 {path.name} ← portal meta")
    return path


def _python_docx_md_to_docx(md_path: Path, docx_path: Path) -> bool:
    try:
        from docx import Document
    except ImportError:
        return False
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            doc.add_paragraph(line)
    doc.save(str(docx_path))
    return True


def _build_pdf_for_md(md_path: Path) -> None:
    pdf_path = md_path.with_suffix(".pdf")
    with tempfile.TemporaryDirectory() as td:
        tmp_docx = Path(td) / "body.docx"
        if not _python_docx_md_to_docx(md_path, tmp_docx):
            print(
                f"경고: {md_path.name} → 중간 docx 실패. pip install python-docx",
                file=sys.stderr,
            )
            return
        try:
            convert_docx_to_pdf(tmp_docx, pdf_path)
        except Exception as e:
            print(f"경고: {md_path.name} → pdf 변환 실패: {e}", file=sys.stderr)
            return
    print(f"PDF {pdf_path.name}")


def main() -> None:
    ap = argparse.ArgumentParser(
        description="commission-portal/public/source MD·PDF·site-display 동기화"
    )
    ap.add_argument(
        "--mirror-public",
        action="store_true",
        help="(호환, 무시) 예전 web/source → public 미러 단계는 제거됨",
    )
    ap.add_argument(
        "--force-site-display",
        action="store_true",
        help="site-display.json 을 무조건 portal meta(없으면 기본값)로 덮어씀",
    )
    ap.add_argument("--no-pdf", action="store_true", help="PDF 생성 생략")
    ap.add_argument(
        "--no-docx",
        action="store_true",
        help="(호환) PDF 생성 생략 — --no-pdf 와 동일",
    )
    args = ap.parse_args()

    if not _PORTAL_DATA.is_file():
        raise SystemExit(f"없음: {_PORTAL_DATA} — 먼저 build_commission_evidence_json.py 실행")

    data = json.loads(_PORTAL_DATA.read_text(encoding="utf-8"))
    meta = data.get("meta") or {}
    tab_sources = meta.get("tabSources") or {}

    _write_site_display(meta, args.force_site_display)
    _copy_tab_sources(tab_sources)

    if not args.no_pdf and not args.no_docx:
        for _k, dest_name in _TAB_FILES:
            md = _PUBLIC_SOURCE / dest_name
            if md.is_file():
                _build_pdf_for_md(md)

    if args.mirror_public:
        print("(참고) --mirror-public 은 더 이상 필요 없습니다. 출력은 항상 public/source 입니다.")


if __name__ == "__main__":
    main()
