# -*- coding: utf-8 -*-
"""행정심판 청구서·집행정지신청서 MD → DOCX → PDF.

- **흐름:** Markdown 정리 → `python-docx`로 DOCX → `pdf_convert_util`(Word/LibreOffice)로 PDF.
- **글꼴(공문 관행):** 본문 기본 **맑은 바탕**(명조·바탕 계열). 행정 공문에서 본문은 고딕보다 **바탕·신명조·맑은 바탕**을 쓰는 경우가 많고, 고딕은 표·화면·강조에 자주 쓰입니다. `--font`로 바꿀 수 있습니다.
- **증거자료 목록:** `## 증거자료 목록` 아래 `1.`, `1-2.` 형태 줄은 **표**(번호 | 갑호증 표시 | 내용)로 넣어 Word에서 직접 수정하기 쉽게 합니다. MD 링크 `[갑 제…](#n)`는 셀에 **표시 문자만** 넣습니다(하이퍼링크는 Word에서 필요 시 추가).

**공문서(행정) 관행 참고값 — 기본 `--layout gongmun`**

- 법령 하나로 전국 단일 수치는 없고, 행안부·자치단체 **표준서식**마다 차이가 있습니다.
- 흔한 관행: **A4(210×297mm)**, 본문 **10.5pt**·바탕 계열, **줄간격 약 160%**(배수 1.6), **여백 상하좌우 약 25mm**(20~30mm·제본용 왼쪽 넓게 등 변형 많음).
- `--layout word` 는 여백·줄간격 미지정·본문 11pt(이전 동작에 가깝게).

필요: `pip install python-docx` 및 Word 또는 LibreOffice(DOCX→PDF).

  python tools/md_submission_to_docx_pdf.py --all
  python tools/md_submission_to_docx_pdf.py -i "행정심판청구(최종)/260405/260405_01_행정심판청구서.md"
  # --all: 청구서·집행정지 + 별지 제1~3호(증거목록·판례주석·시간축) MD가 있으면 함께 DOCX 생성
  python tools/md_submission_to_docx_pdf.py --all --no-pdf   # DOCX만
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

_REPO = Path(__file__).resolve().parent.parent
_DEFAULT_DIR = _REPO / "행정심판청구(최종)" / "260405"

# 공문 관행 참고(기관 표준서식과 다를 수 있음)
_GONGMUN_MARGIN_MM = 25.0
_GONGMUN_BODY_PT = 10.5
_GONGMUN_LINE_MULT = 1.6
_GONGMUN_TABLE_PT = 10.0

_WORD_DEFAULT_BODY_PT = 11.0
_WORD_DEFAULT_TABLE_PT = 10.0

_EVIDENCE_LINE = re.compile(r"^(\d+(?:-\d+)?)\.\s+(.+)$")
_CONTINUATION = re.compile(r"^\s{2,}(?:—|-.+)")


def _html_inline_to_md_frag(html: str) -> str:
    """`<strong>…</strong>` 등을 `**…**` 로 바꿔 DOCX 파서가 굵게 처리하도록 함."""
    s = html
    s = re.sub(r"<strong>(.+?)</strong>", r"**\1**", s, flags=re.DOTALL | re.IGNORECASE)
    s = re.sub(r"<b>(.+?)</b>", r"**\1**", s, flags=re.DOTALL | re.IGNORECASE)
    return s


def _strip_md_wrapper(raw: str) -> str:
    s = re.sub(r"<style[^>]*>.*?</style>", "", raw, flags=re.DOTALL | re.IGNORECASE)
    s = re.sub(r"<!--.*?-->", "", s, flags=re.DOTALL)
    s = re.sub(r'<div\s+class="doc-gongmun"\s*>', "", s, flags=re.IGNORECASE)
    # 미리보기용 HTML 수신·서명 → DOCX 우측 정렬: 수신 `>> `, 서명 `>>! `(본문 pt×1.45)
    s = re.sub(
        r'<p\s+class="hdr-recipient"[^>]*>\s*(.+?)\s*</p>',
        lambda m: ">> "
        + re.sub(r"\s+", " ", _html_inline_to_md_frag(m.group(1).strip())),
        s,
        flags=re.DOTALL | re.IGNORECASE,
    )
    # 서명(청구인·신청인): 귀중과 동일 배율로 DOCX에서도 키우기 → `>>! ` (md_to_docx에서 처리)
    s = re.sub(
        r'<p\s+class="ftr-sign"[^>]*>\s*(.+?)\s*</p>',
        lambda m: ">>! "
        + re.sub(r"\s+", " ", _html_inline_to_md_frag(m.group(1).strip())),
        s,
        flags=re.DOTALL | re.IGNORECASE,
    )
    s = s.rstrip()
    if s.lower().endswith("</div>"):
        s = s[: -len("</div>")].rstrip()
    return s.strip() + "\n"


def _plain_from_md_links(text: str) -> str:
    """[표시](url) → 표시, <url> 유지."""

    def repl(m: re.Match[str]) -> str:
        return m.group(1)

    return re.sub(r"\[([^\]]+)\]\([^)]+\)", repl, text)


def _set_section_margins_mm(doc, mm: float) -> None:
    from docx.shared import Mm

    m = Mm(mm)
    for sec in doc.sections:
        sec.top_margin = m
        sec.bottom_margin = m
        sec.left_margin = m
        sec.right_margin = m


def _set_list_and_normal_spacing(
    doc, *, line_multiple: float, body_pt: float, body_font: str
) -> None:
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for name in ("Normal", "List Bullet", "List Number"):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        pf = st.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_multiple
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        st.font.size = Pt(body_pt)
        st.font.name = body_font
        rPr = st.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), body_font)
        rFonts.set(qn("w:ascii"), body_font)
        rFonts.set(qn("w:hAnsi"), body_font)


def _set_doc_fonts(doc, body_font: str, *, body_pt: float) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = body_font
        if name == "Normal":
            st.font.size = Pt(body_pt)
        elif name == "Heading 1":
            st.font.size = Pt(16)
        elif name == "Heading 2":
            st.font.size = Pt(14)
        elif name == "Heading 3":
            st.font.size = Pt(12)
        else:
            st.font.size = Pt(12)
        rPr = st.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), body_font)
        rFonts.set(qn("w:ascii"), body_font)
        rFonts.set(qn("w:hAnsi"), body_font)


def _justify_paragraph(p) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _center_paragraph(p) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _right_paragraph(p) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# 갑호증 표시·법령 인용 등: 본문과 구분만 되게 은은한 남색(Word RGB)
_EVIDENCE_RGB = (0x3D, 0x5A, 0x80)
_GABHO_RE = re.compile(r"(갑 제[\d\-]+호증)")


def _add_runs_with_style(
    p,
    text: str,
    body_pt: float,
    *,
    evidence_blue: bool = True,
) -> None:
    """`**굵게**` 및 (선택) `갑 제N호증` 문구에 스타일 적용. 링크는 먼저 제거한 문자열 기준."""
    from docx.shared import Pt

    plain = _plain_from_md_links(text)
    pos = 0
    while pos < len(plain):
        m_bold = re.search(r"\*\*(.+?)\*\*", plain[pos:])
        next_bold = m_bold.start() + pos if m_bold else len(plain)
        seg = plain[pos:next_bold]
        if seg:
            _add_runs_gabho_colored(p, seg, body_pt, evidence_blue=evidence_blue)
        if m_bold:
            r = p.add_run(m_bold.group(1))
            r.bold = True
            r.font.size = Pt(body_pt)
            pos = next_bold + len(m_bold.group(0))
        else:
            break


def _add_runs_gabho_colored(
    p,
    seg: str,
    body_pt: float,
    *,
    evidence_blue: bool,
) -> None:
    from docx.shared import Pt, RGBColor

    if not evidence_blue:
        r = p.add_run(seg)
        r.font.size = Pt(body_pt)
        return
    last = 0
    for m in _GABHO_RE.finditer(seg):
        if m.start() > last:
            r = p.add_run(seg[last : m.start()])
            r.font.size = Pt(body_pt)
        r = p.add_run(m.group(1))
        r.font.size = Pt(body_pt)
        r.font.color.rgb = RGBColor(*_EVIDENCE_RGB)
        last = m.end()
    if last < len(seg):
        r = p.add_run(seg[last:])
        r.font.size = Pt(body_pt)


def _add_paragraph(
    doc,
    text: str,
    *,
    body_pt: float,
    style: str | None = None,
    justify: bool = True,
):
    from docx.shared import Pt

    p = doc.add_paragraph(style=style)
    if text:
        _add_runs_with_style(p, text, body_pt, evidence_blue=True)
    if justify:
        _justify_paragraph(p)
    return p


def _add_right_paragraph(doc, text: str, *, body_pt: float) -> None:
    p = doc.add_paragraph()
    _add_runs_with_style(p, text, body_pt, evidence_blue=True)
    _right_paragraph(p)


def _add_spacer_paragraph(doc, *, body_pt: float, before: float = 0, after: float = 5) -> None:
    """빈 줄·구분선에 대응하는 여백 단락."""
    from docx.shared import Pt

    p = doc.add_paragraph()
    r = p.add_run("\u200b")
    r.font.size = Pt(1)
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def _is_hr_line(stripped: str) -> bool:
    return bool(re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", stripped))


def _add_table_evidence(
    doc,
    rows: list[tuple[str, str, str]],
    body_font: str,
    *,
    table_pt: float,
) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    if not rows:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = ("목록번호", "갑호증(표시)", "내용·비고")
    for c, h in enumerate(hdr):
        cell = tbl.rows[0].cells[c]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(table_pt)
                run.font.name = body_font
                r = run._element.get_or_add_rPr()
                rf = r.get_or_add_rFonts()
                rf.set(qn("w:eastAsia"), body_font)
    for ri, (num, label, desc) in enumerate(rows, start=1):
        cells = tbl.rows[ri].cells
        for ci, val in enumerate((num, label, desc)):
            cells[ci].text = val
            for paragraph in cells[ci].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci < 2 else WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in paragraph.runs:
                    run.font.size = Pt(table_pt)
                    run.font.name = body_font
                    r = run._element.get_or_add_rPr()
                    rf = r.get_or_add_rFonts()
                    rf.set(qn("w:eastAsia"), body_font)
    doc.add_paragraph()


def _parse_evidence_label_desc(rest: str) -> tuple[str, str]:
    """`[갑 제1-1호증](#1-1)(설명...)` → (갑 제1-1호증, 설명...)."""
    rest = rest.strip()
    m = re.match(r"^\[([^\]]+)\]\([^)]+\)\s*(.*)$", rest)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return _plain_from_md_links(rest), ""


def md_to_docx(
    md_path: Path,
    docx_path: Path,
    *,
    body_font: str,
    layout: str = "gongmun",
) -> None:
    try:
        from docx import Document
    except ImportError as e:
        raise SystemExit("python-docx 필요: pip install python-docx") from e
    from docx.shared import Pt

    if layout == "gongmun":
        body_pt = _GONGMUN_BODY_PT
        table_pt = _GONGMUN_TABLE_PT
        margin_mm = _GONGMUN_MARGIN_MM
        line_mult = _GONGMUN_LINE_MULT
    else:
        body_pt = _WORD_DEFAULT_BODY_PT
        table_pt = _WORD_DEFAULT_TABLE_PT
        margin_mm = None
        line_mult = None

    raw = md_path.read_text(encoding="utf-8")
    text = _strip_md_wrapper(raw)
    lines = text.splitlines()

    doc = Document()
    _set_doc_fonts(doc, body_font, body_pt=body_pt)
    if margin_mm is not None and line_mult is not None:
        _set_section_margins_mm(doc, margin_mm)
        _set_list_and_normal_spacing(
            doc, line_multiple=line_mult, body_pt=body_pt, body_font=body_font
        )

    i = 0
    n = len(lines)
    evidence_intro = False  # 「증거자료 목록」 직후 ~ 첫 번호 줄 전
    evidence_table_active = False
    evidence_rows: list[tuple[str, str, str]] = []
    pending_cont: list[str] = []

    def flush_evidence_table() -> None:
        nonlocal evidence_rows, evidence_table_active
        if evidence_rows:
            _add_table_evidence(
                doc, evidence_rows, body_font, table_pt=table_pt
            )
        evidence_rows = []
        evidence_table_active = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if re.match(r"<div[^>]*break-before:\s*page", line, re.I) or (
            stripped.lower().startswith("<div") and "page" in stripped.lower()
        ):
            flush_evidence_table()
            doc.add_page_break()
            i += 1
            continue

        if stripped == "":
            if not evidence_table_active:
                _add_spacer_paragraph(doc, body_pt=body_pt, before=0, after=5)
            i += 1
            continue

        if _is_hr_line(stripped):
            if not evidence_table_active:
                _add_spacer_paragraph(doc, body_pt=body_pt, before=6, after=10)
            i += 1
            continue

        if stripped.startswith(">>! "):
            flush_evidence_table()
            _add_right_paragraph(
                doc, stripped[4:].strip(), body_pt=round(body_pt * 1.45, 2)
            )
            i += 1
            continue

        if stripped.startswith(">> "):
            flush_evidence_table()
            _add_right_paragraph(doc, stripped[3:].strip(), body_pt=body_pt)
            i += 1
            continue

        if stripped.startswith("# "):
            flush_evidence_table()
            doc.add_heading(_plain_from_md_links(stripped[2:].strip()), level=1)
            if doc.paragraphs:
                _center_paragraph(doc.paragraphs[-1])
            i += 1
            continue
        if stripped.startswith("## "):
            flush_evidence_table()
            evidence_intro = False
            h = stripped[3:].strip()
            doc.add_heading(_plain_from_md_links(h), level=2)
            if h == "증거자료 목록":
                evidence_intro = True
            i += 1
            continue
        if stripped.startswith("### "):
            flush_evidence_table()
            doc.add_heading(_plain_from_md_links(stripped[4:].strip()), level=3)
            i += 1
            continue
        if stripped.startswith("#### "):
            flush_evidence_table()
            doc.add_heading(_plain_from_md_links(stripped[5:].strip()), level=4)
            i += 1
            continue

        if evidence_intro or evidence_table_active:
            em = _EVIDENCE_LINE.match(stripped)
            if em:
                if evidence_intro:
                    evidence_intro = False
                    evidence_table_active = True
                if pending_cont and evidence_rows:
                    a, b, c = evidence_rows[-1]
                    evidence_rows[-1] = (a, b, c + " " + " ".join(pending_cont))
                    pending_cont = []
                num, rest = em.group(1), em.group(2)
                label, desc = _parse_evidence_label_desc(rest)
                evidence_rows.append((num, label, desc))
                i += 1
                continue
            if evidence_table_active and _CONTINUATION.match(line) and evidence_rows:
                pending_cont.append(_plain_from_md_links(stripped))
                i += 1
                continue
            if evidence_table_active:
                if pending_cont and evidence_rows:
                    a, b, c = evidence_rows[-1]
                    evidence_rows[-1] = (a, b, c + " " + " ".join(pending_cont))
                    pending_cont = []
                flush_evidence_table()
                evidence_intro = False
                # 현재 줄을 일반 본문으로 다시 처리
                continue
            if evidence_intro:
                if stripped.startswith("* ") or stripped.startswith("- "):
                    p = doc.add_paragraph(style="List Bullet")
                    _add_runs_with_style(
                        p, stripped[2:].strip(), body_pt, evidence_blue=True
                    )
                    _justify_paragraph(p)
                else:
                    _add_paragraph(doc, stripped, body_pt=body_pt)
                i += 1
                continue

        if stripped.startswith("* ") or stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_style(p, stripped[2:].strip(), body_pt, evidence_blue=True)
            _justify_paragraph(p)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped) and not evidence_intro and not evidence_table_active:
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_style(
                p, re.sub(r"^\d+\.\s+", "", stripped), body_pt, evidence_blue=True
            )
            _justify_paragraph(p)
            i += 1
            continue

        _add_paragraph(doc, stripped, body_pt=body_pt)
        i += 1

    if pending_cont and evidence_rows:
        a, b, c = evidence_rows[-1]
        evidence_rows[-1] = (a, b, c + " " + " ".join(pending_cont))
    flush_evidence_table()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def convert_one(
    md_path: Path,
    *,
    body_font: str,
    layout: str,
    docx_path: Path | None,
    pdf_path: Path | None,
    write_pdf: bool,
) -> tuple[Path, Path | None]:
    md_path = md_path.resolve()
    out_docx = docx_path or md_path.with_suffix(".docx")
    out_pdf = pdf_path or md_path.with_suffix(".pdf")
    if not out_docx.is_absolute():
        out_docx = (_REPO / out_docx).resolve()
    if out_pdf is not None and not out_pdf.is_absolute():
        out_pdf = (_REPO / out_pdf).resolve()

    md_to_docx(md_path, out_docx, body_font=body_font, layout=layout)
    pdf_done = None
    if write_pdf:
        from pdf_convert_util import convert_docx_to_pdf

        convert_docx_to_pdf(out_docx, out_pdf)
        pdf_done = out_pdf
    return out_docx, pdf_done


def main() -> int:
    ap = argparse.ArgumentParser(description="제출용 MD → DOCX [→ PDF]")
    ap.add_argument("-i", "--input", type=Path, default=None)
    ap.add_argument("-o", "--output", type=Path, default=None, help="DOCX 경로(단일 입력)")
    ap.add_argument("--pdf", type=Path, default=None, help="PDF 경로(기본: md와 동일 이름 .pdf)")
    ap.add_argument("--no-pdf", action="store_true", help="DOCX만 생성")
    ap.add_argument("--dir", type=Path, default=_DEFAULT_DIR)
    ap.add_argument("--all", action="store_true")
    ap.add_argument(
        "--font",
        default="맑은 바탕",
        help="본문·표 기본 글꼴(기본: 맑은 바탕 — 공문 본문 관행). 예: 바탕, 신명조, 맑은 고딕",
    )
    ap.add_argument(
        "--layout",
        choices=("gongmun", "word"),
        default="gongmun",
        help="gongmun: 여백 25mm·본문 10.5pt·줄간격 1.6배(관행 참고). word: 여백·줄간격 미설정·11pt.",
    )
    args = ap.parse_args()

    write_pdf = not args.no_pdf

    if args.all:
        base = args.dir.resolve()
        _batch = (
            "260405_01_행정심판청구서.md",
            "260405_02_집행정지신청서.md",
            "260405_별지제1호_증거자료_목록.md",
            "260405_별지제2호_주요인용판례_및_적용주석.md",
            "260405_별지제3호_사실관계_시간축_정리표.md",
        )
        for name in _batch:
            md = base / name
            if not md.is_file():
                print(f"건너뜀: {md}", file=sys.stderr)
                continue
            try:
                dx, pdf = convert_one(
                    md,
                    body_font=args.font,
                    layout=args.layout,
                    docx_path=None,
                    pdf_path=None,
                    write_pdf=write_pdf,
                )
                print(f"DOCX: {dx.relative_to(_REPO)}")
                if pdf:
                    print(f"PDF:  {pdf.relative_to(_REPO)}")
            except Exception as e:
                print(f"오류 {name}: {e}", file=sys.stderr)
                return 1
        return 0

    if args.input is None:
        print("-i 또는 --all", file=sys.stderr)
        return 1
    md_in = args.input
    if not md_in.is_absolute():
        md_in = (_REPO / md_in).resolve()
    out_docx = args.output
    out_pdf = args.pdf
    if out_docx is not None and not out_docx.is_absolute():
        out_docx = (_REPO / out_docx).resolve()

    try:
        dx, pdf = convert_one(
            md_in,
            body_font=args.font,
            layout=args.layout,
            docx_path=out_docx,
            pdf_path=out_pdf,
            write_pdf=write_pdf,
        )
        print(f"DOCX: {dx.relative_to(_REPO)}")
        if pdf:
            print(f"PDF:  {pdf.relative_to(_REPO)}")
    except Exception as e:
        print(str(e), file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
